import os
import requests
from docx import Document
import tkinter as tk
from tkinter import simpledialog

def download_file(url, output_path):
    response = requests.get(url)
    if response.status_code == 200:
        with open(output_path, 'wb') as file:
            file.write(response.content)
        print(f"Downloaded {url} to {output_path}")
    else:
        print(f"Failed to download {url}. Status code: {response.status_code}")

def get_files_from_repo(user, repo, branch='main'):
    api_url = f"https://api.github.com/repos/{user}/{repo}/git/trees/{branch}?recursive=1"
    response = requests.get(api_url)
    if response.status_code == 200:
        tree = response.json().get('tree', [])
        files = [item['path'] for item in tree]
        return files
    else:
        print(f"Failed to fetch repository tree. Status code: {response.status_code}")
        return []

def download_selected_files(user, repo, output_dir, branch='main', selected_indices=None):
    files = get_files_from_repo(user, repo, branch)
    if not files:
        print("No files found in the repository.")
        return

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    document = Document()

    file_types = {}
    for file_path in files:
        file_type = file_path.split('.')[-1].lower()
        if file_type not in file_types:
            file_types[file_type] = []
        file_types[file_type].append(file_path)

    file_type_selection = simpledialog.askstring("Input", f"Select file types to include in the Word document (separated by commas):\n{', '.join(file_types.keys())}")

    if not file_type_selection:
        print("No file types selected. Exiting...")
        return

    selected_files = []
    selected_file_paths = []
    for file_type in file_type_selection.split(","):
        if file_type.lower() in file_types:
            selected_files.extend(file_types[file_type.lower()])
            selected_file_paths.extend([file_path for file_path in file_types[file_type.lower()]])
        else:
            print(f"No files of type {file_type} found.")

    if not selected_files:
        print("No files selected. Exiting...")
        return

    print("Available files:")
    for idx, file_path in enumerate(selected_file_paths, start=1):
        print(f"{idx}. {file_path}")

    selected_indices = input("Enter the indices of files you want to include in the Word document (separated by commas): ").split(",")
    selected_files = [selected_files[int(idx) - 1] for idx in selected_indices]

    first_file = True
    for file_path in selected_files:
        raw_url = f"https://raw.githubusercontent.com/{user}/{repo}/{branch}/{file_path}"
        response = requests.get(raw_url)
        if response.status_code == 200:
            content = response.text
            if first_file:
                first_file = False
            else:
                document.add_page_break()
            document.add_heading(file_path, level=1)
            document.add_paragraph(content)
            print(f"Added content of {file_path} to Word document")
        else:
            print(f"Failed to download {file_path}. Status code: {response.status_code}")

    docx_path = os.path.join(output_dir, f"{repo}_files.docx")
    document.save(docx_path)
    print(f"Word document saved to: {docx_path}")

def get_input():
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    user = simpledialog.askstring("Input", "Enter your GitHub username:")
    repo = simpledialog.askstring("Input", "Enter the name of the repository:")
    output_dir = "github_files" # Replace with your desired output directory

    download_selected_files(user, repo, output_dir)

if __name__ == "__main__":
    get_input()
